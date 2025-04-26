# written by <Linsur>
# please run this code in terminal or command line to get email attachments and save them to docx files.

import poplib
import email
import os
import email.utils
from docx import Document
from bs4 import BeautifulSoup
from docx.oxml.ns import qn 
import traceback

def decode(header):
    if header is None:
        return "无标题"
    decoded_parts = email.header.decode_header(header)
    subject = ""
    for decoded_part in decoded_parts:
        if isinstance(decoded_part[0], bytes):
            subject += decoded_part[0].decode(decoded_part[1] or 'utf-8')
        else:
            subject += decoded_part[0]
    return subject

"""
    将邮件正文粘贴到word文档
"""

def clean_filename(filename):
    return ''.join(c if c.isalnum() or c in '._-' else '_' for c in filename)

def save_email_body_to_docx(msg, save_dir):
    subject = decode(msg.get('Subject'))
    cleaned_subject = clean_filename(subject)
    docx_file_path = os.path.join(save_dir, f"{cleaned_subject}.docx")

    doc = Document()

    for part in msg.walk():
        content_type = part.get_content_type()
        content_disposition = part.get_content_disposition()

        if content_disposition != 'attachment':
            if content_type == 'text/plain':
                body = part.get_payload(decode=True).decode(part.get_content_charset() or 'utf-8')
                doc.add_paragraph(body)
            elif content_type == 'text/html':
                html_body = part.get_payload(decode=True).decode(part.get_content_charset() or 'utf-8')
                soup = BeautifulSoup(html_body, 'html.parser')
                for element in soup.find_all(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    if element.name == 'p':
                        doc.add_paragraph(element.get_text())
                    elif element.name == 'div':
                        doc.add_paragraph(element.get_text())
                    elif element.name.startswith('h'):
                        doc.add_heading(element.get_text(), level=int(element.name[1]))

    try:
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.save(docx_file_path)
        print(f"Email body saved to docx: {docx_file_path}")
    except Exception as e:
        print(f"Failed to save email body to docx {cleaned_subject}: {e}")
        traceback.print_exc()

"""
    下载邮箱附件
"""

def download_attachment(msg, save_dir):
    for part in msg.walk():
        if part.get_content_disposition() == 'attachment':
            attachment_name = email.utils.unquote(decode(part.get_filename()))
            if not attachment_name:
                print("Attachment name is empty, skipping...")
                continue

            attachment_name = os.path.basename(attachment_name)
            attachment_content = part.get_payload(decode=True)
            if not attachment_content:
                print("Attachment content is empty, skipping...")
                continue

            attachment_file_path = os.path.join(save_dir, attachment_name)

            try:
                with open(attachment_file_path, 'wb') as attachment_file:
                    attachment_file.write(attachment_content)
                print(f"Attachment saved: {attachment_file_path}")
            except Exception as e:
                print(f"Failed to save attachment {attachment_name}: {e}")

def main():
    try:
        server = poplib.POP3(host='pop.126.com')
        print("Connected to POP3 server")
    except Exception as e:
        print(f"Failed to connect to POP3 server: {e}")
        return

    try:
        server.user('xxxxxxxxx@126.com')
        server.pass_('xxxxxxxxxx')
        print("Login successful")
    except Exception as e:
        print(f"Login failed: {e}")
        server.quit()
        return

    try:
        msg_count, _ = server.stat()
    except Exception as e:
        print(f"Failed to get email count: {e}")
        server.quit()
        return

    save_dir = "D:\\email_attachments"
    if not os.path.exists(save_dir):
        try:
            os.makedirs(save_dir)
            print(f"Created directory: {save_dir}")
        except Exception as e:
            print(f"Failed to create directory {save_dir}: {e}")
            server.quit()
            return

    for i in range(msg_count):
        try:
            _, lines, _ = server.retr(i + 1)
        except Exception as e:
            print(f"Failed to retrieve email {i + 1}: {e}")
            continue

        msg_bytes_content = b'\r\n'.join(lines)
        try:
            msg = email.message_from_bytes(msg_bytes_content)
        except Exception as e:
            print(f"Failed to parse email {i + 1}: {e}")
            continue

        save_email_body_to_docx(msg, save_dir)
        download_attachment(msg, save_dir)

    server.quit()
    print("Disconnected from POP3 server")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"Error in main: {e}")
        traceback.print_exc()
